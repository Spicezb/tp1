# Trabajo realizado Por Luis Guillermo Alfaro Chacón y Xavier Céspedes Alvarado.
# Fecha de inicio: 23/04/2025 a las 12:00        
#
# Versión de python: 3.13.2
# Se debe instalar las librerías names, fpdf y python-docx

# Importación de librerias
import names 
import re
import random
import pickle
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from respaldarEnXML import *
from aplazados import *
from generarCurva import *
from docx import Document
from datetime import datetime

def notas():
    """
    Funcionamiento:
    - Es la auxiliar que solicita los porcentajes de cada evaluación.
    Entradas:
    N/A
    Salidas:
    Se retornan los 3 porcentajes correspondientes.
    """
    while True:
        try:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            porce1=int(input("Indique el porcentaje de la primer evaluación: "))
            porce2=int(input("Indique el porcentaje de la segundo evaluación: "))
            porce3=int(input("Indique el porcentaje de la tercer evaluación: "))
            if porce1+porce2+porce3 != 100:
                raise TypeError
            return porce1,porce2,porce3
        except TypeError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("El porcentaje Total de las 3 evaluaciones debe de ser igual a 100%.\n" \
            f"La suma de los porcentajes ingresados es de: {str(porce1+porce2+porce3)}.\n"
            "Por favor, Ingrese nuevamente los porcentajes.")
            input("Presione enter para continuar.")
            os.system("cls")
        except ValueError:
            print("******************** Crear Base de Datos ********************")
            os.system("cls")
            print("Debe ingresar valores válidos. Por ejemplo:\n" \
            "Indique el porcentaje de la primer evaluación: 35")
            input("Presione enter para continuar.")
            os.system("cls")

def crearNotas(porce1,porce2,porce3):
    """
    Funcionamiento:
    - Crea las notas de manera aleatoria mayores a 0 y menores a 101.
    Entradas:
    - porce1 (int): contiene el porcentaje de la primer evaluación.
    - porce2 (int): contiene el porcentaje de la segunda evaluación.
    - porce3 (int): contiene el porcentaje de la tercer evaluación.
    Salidas:
    Se retorna una tupla con las 3 notas aleatorias, la nota final del curso y la nota de curva.
    """
    evalu1=random.randint(1, 100)
    evalu2=random.randint(1, 100)                                                   # Genera numeros aleatorios del 1 al 100.
    evalu3=random.randint(1, 100)
    notaFinal=round((evalu1*porce1/100)+(evalu2*porce2/100)+(evalu3*porce3/100),2)  # Operación para generar la nota final.
    tupla=(evalu1,evalu2,evalu3,notaFinal,notaFinal)
    return tupla

def crearCorreo(nombre,carne):
    correo=nombre[0][:1]+nombre[1]+carne[6:]+"@estudiantec.cr"
    correo=correo.lower()
    return correo.replace("á","a").replace("é","e").replace("í","i").replace("ó","o").replace("ú","u")

def crearCarne(ini, fin):
    lista=[]
    txtSedes=open("sedes.txt","r")         
    x=txtSedes.readlines()
    for i in range(len(x)):
        if i+1<10:
            lista.append("0"+str(i+1))
        else:
            lista.append(str(i+1))          
    rand=str(random.randint(0000,9999))
    while not re.match(r"\d{4}",rand):
        rand="0"+rand
    carnet=str(random.randint(ini,fin))+random.choice(lista)+rand
    txtSedes.close()
    return str(carnet)

def crearNombres():
    """
    Funcionamiento:
    Se crean nombres de manera aleatoria, gracias a la libreia names y random, donde se le asigna un genero al nombre,
    y se crea el nombre dependiendo del genero dado. Luego la informacion se almacena por separado.
    Entradas:
    N/A
    Salidas:
    Se retorna una lista con una tupla que contiene el nombre y los apellidos, y el genero por aparte, este de manera Booleana
    """
    nombrePersona=""                                    # Se crea la variable que contendra el nombre creado.
    genero=random.choice(("male","female"))             # Se pide que escoja un genero para que cree los nombres basados a este género.
    persona=names.get_first_name(gender=genero)         # Se crea el nombre.
    primerApe=names.get_last_name()                     
    segundoApe=names.get_last_name()                    # Se crean los 2 apellifos
    if genero=="male":
        genero = True
    else:
        genero= False                                   # Al tener los generos en ingles, se pasan booleano para trabajar mejor.
    nombrePersona = f"{persona},{primerApe},{segundoApe},{genero}\n"    # Crea la información completa del nombre.
    return nombrePersona

def llenarBD(nomb,rango,rango2,notas):
    """
    Funcionamiento:
    - Se encarga de llenar la información de cada persona y retornala para ser agregada en la base de datos. 
    Entradas:
    - nomb(tuple): contiene el nombre, apellidos y genero de cada persona.
    - rango(int): contiene el año incial de los carnés.
    - rango2(int): contiene el año final de los carnés.
    - notas(tuple): contiene las notas aleatorias creadas por el sistema.
    Salidas:
    - Retorna la lista con la información de cada persona.
    """
    infoPerso=[]                                        # Lista con la información
    nombre=(nomb[0],nomb[1],nomb[2])                    # Se desgloza la tupla, para obtener el nombre y el apellido de la persona
    genero=nomb[3]                                      # Se obtiene el genero
    if genero == "Masculino" or genero == "True":
        genero=True                                     # Como la libreía names crea los generos en ingles, se igualan todos los generos a un valor Booleano.
    else:
        genero=False
    carne= crearCarne(rango,rango2)
    correo = crearCorreo(nombre,str(carne))         
    infoPerso.append(nombre)
    infoPerso.append(genero)                            # Se agregan todos los valores a la lista que llena la Base de datos.
    infoPerso.append(carne)
    infoPerso.append(correo)
    infoPerso.append(notas)
    return(infoPerso)

def cantidadEstu():
    """
    Funcionamiento:
    - Es la auxiliar que solicita la cantidad de estudiantes a crear y el porcentaje a utilizar.
    Entradas:
    N/A
    Salidas:
    Se retorna una lista con una tupla los 2 valores correspondientes.
    """
    conta=1
    while True:
        try:
            print("******************** Crear Base de Datos ********************")
            cantidadCrear=int(input("Digite la cantidad de estudiantes a crear: "))
            break
        except ValueError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("Debe de ingresar una cantidad válida.\nPor ejemplo:\n" \
            "Digite la cantidad de estudiantes a crear: 120\n" \
            "Digite el porcentaje a agregar de las fuentes: 5\n")
            input("Presione enter para continuar.")
            os.system("cls")
    while True:
        try:
            if conta != 1:
                print("******************** Crear Base de Datos ********************")
                print("Digite la cantidad de estudiantes a crear: "+str(cantidadCrear))
                porcentaje=int(input("Digite el porcentaje a agregar de las fuentes: "))
            else:
                conta+=1
                porcentaje=int(input("Digite el porcentaje a agregar de las fuentes: "))
            break
        except ValueError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("Debe de ingresar una cantidad válida.\nPor ejemplo:\n" \
            "Digite la cantidad de estudiantes a crear: 120\n" \
            "Digite el porcentaje a agregar de las fuentes: 5\n")
            input("Presione enter para continuar.")
            os.system("cls")
    os.system("cls")
    return (cantidadCrear,porcentaje)

def generaciones():   
    """
    Funcionamiento:
    - Es la auxiliar que solicita el rango de años para generar los carnés.
    Entradas:
    N/A
    Salidas:
    Se retorna una lista con una tupla los 2 años correspondientes.
    """
    conta=1
    while True:
        try:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("Ingrese el rango de años para generar los carnets")
            rango=int(input("Año inicial: "))
            if not 999<rango<10000:
                raise ValueError
            break
        except ValueError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("El año no es válido, debe ingresar un año de 4 dígitos.\nPor ejemplo: 2022")
            input("Presione enter para continuar.")
            os.system("cls")
    while True:
        try:
            if conta != 1:
                os.system("cls")
                print("******************** Crear Base de Datos ********************")
                print("Ingrese el rango de años para generar los carnets")
                print("Año inicial: "+str(rango))
                rango2=int(input("Año final: "))
            else:
                conta+=1
                rango2=int(input("Año final: "))
            
            if rango>rango2:
                raise TypeError
            elif rango2>9999:
                raise ValueError
            break
        except ValueError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("El año no es válido, debe ingresar un año de 4 dígitos.\nPor ejemplo: 2024")
            input("Presione enter para continuar.")
            os.system("cls")
        except TypeError:
            os.system("cls")
            print("******************** Crear Base de Datos ********************")
            print("El año final no puede ser menor que el inicial.\nIngrese por ejemplo:\n" \
            "Año inicial: 2022\nAño final: 2024")
            input("Presione enter para continuar.")
            os.system("cls")
    return (rango,rango2)

def crearBD(archivo,porce1,porce2,porce3):
    """
    Funcionamiento:
    - Crea la base de Datos donde se va a estar almacenando la información del programa.
    Entradas:
    - archivo(str): contiene la información toda la información de la Base de Datos.
    - porce1: contiene el primer porcentaje de las evaluaciones.
    - porce2: contiene el primer porcentaje de las evaluaciones.
    - porce3: contiene el primer porcentaje de las evaluaciones.
    Salidas:
    - Retorna un True al realizar el proceso correctamente, pero este no se usa ni se ve en ninguna parte. Antes de el True, 
    si se le muestra al usuario que la Base de datos se creo correctamente.
    """
    # Definición de variables
    lista=[]                                        # Lista que va a almacenar la información.
    lstnombsconv=[]                                 # Lista que contiene los nombres de las 2 fuentes distintas.
    abrirAr=open(archivo,"wb")                      # Crea el archivo binario.
    cantiDeEstu=cantidadEstu()                      # Contiene la cantidad de estudiantes, dada por la funcion cantidadEstu.
    rangos=generaciones()                           # Contiene los rangos de las generaciones a crear.
    txtNombresGenerados=open("nombres.txt","w")     # Crea el archivo de nombres generados por el programa.
    # Se le crean los nombres aleatorios.
    for i in range(cantiDeEstu[0]):
        txtNombresGenerados.write(crearNombres())
    txtNombresGenerados.close()
    # Se da la cantidad de estudiantes a agregar de la fuente que nosotros creamos.
    nombresArchivo=escogerDeArchi("nombres.txt",cantiDeEstu[1],lstnombsconv)                                  # Se agregan los nombres creados por nosotros.
    nombresArchivo.extend(escogerDeArchi("estudiantes.txt",cantiDeEstu[1],lstnombsconv))                      # Se agregan los nombres de la fuente de estudiantes, dada por la profe.
    # Se agregan todos los elementos a la lista de la base de datos.
    for i in range(len(nombresArchivo)):
        lista.append(llenarBD(nombresArchivo[i],rangos[0],rangos[1], crearNotas(porce1,porce2,porce3)))
    pickle.dump(lista,abrirAr)
    abrirAr.close()
    os.system("cls")
    print("******************** Crear Base de Datos ********************")
    print("Base de datos creada y llenada exitosamente.")
    return True

def porcentajeDeEstudiantes(nombreArchivo,porcentaje):
    """
    Funcionamiento:
    - Indica el porcentaje de nombres que se debe de escoger de cada archivo.
    Entradas:
    - porcentaje (int): contiene la cantidad de porcentaje de estudiantes que se va a utilizar en las 2 fuentes.
    - nombreArchivo (str): contiene el nombre del archivo a usar. 
    Salidas:
    - Retorna el porcentaje real y redondeado de la cantidad de estudiantes a agregar dependiendo la fuente.
    """
    conta=0                                                                         # Lleva el conteo de las lineas del archivo.
    archivo=open(nombreArchivo, "r", encoding="utf-8")                              # Se abre el archivo correspondiente.
    for i in archivo:
        conta+=1                                                                    # Se obtiene la cantidad de lineas del archivo.
    if ((porcentaje/100)*conta)%1 >= 0.5:
        conta=((porcentaje/100)*conta)+(((porcentaje/100)*conta)%1)-1               # Operaciones para sacar el porcentaje.
    else:
        conta=((porcentaje/100)*conta)-((porcentaje/100)*conta)%1
    archivo.close()
    return conta

def escogerDeArchi(nombreArchivo,porcentaje,lis):
    """
    Funcionamiento:
    - Escoge la cantidad de nombres según el porcentaje dado, desde el archivo correspondiente.
    Entradas:
    - nombreArchivo (str): contiene el nombre del archivo.
    - porcentaje (int): contiene el porcentaje de estudiantes que se va a utilizar en las 2 fuentes.
    - lis (list): lista que contiene información y la agrega a la base de datos. 
    Salidas:
    - Retorna la lista de los nombres que se escogieron aleatoriamente.
    """
    conta=porcentajeDeEstudiantes(nombreArchivo, porcentaje)                # Contiene la cantidad de estudiantes que hay que escoger.
    estudi=open(nombreArchivo,"r",encoding="utf-8")                         # Abre el contenido del archivo donde hay que sacar los nombres.
    lineas=estudi.readlines()
    estudiante=random.sample(lineas,int(conta))                             # Se saca un estudiante random, sin posibilidad de repetirse.(Esto por el random.sample, que escoge un valor
    for i in range(int(conta)):                                             # random único entre todas las lineas, n cantidad de veces, en este caso escoge conta cantidad de nombres unicos)
        tupla=tuple(estudiante[i].strip().split(","))                       # Se agrega uno por uno los estudiantes escogidos de forma aleatoria, el strip es para limpiar caracteres y   
        lis.append(tupla)                                                   # el split para separar las palabras cuando exite una coma.
    estudi.close()
    return lis

def agregarEstudiante(archivo,estudiante):
    """
    Funcionamiento: 
    - Toma una lista con los datos del estudiante y la agrega a otra lista que contiene a todos los otros estudiantes, esta última es agregada a la base de datos que se vuelve a crear pero actualizada.
    Entradas:
    - archivo(str): Es el archivo de la base de datos.
    - estudiante(list): Contiene los datos del estudiante que se va a agregar.
    Salidas: 
    - Retorna un texto que dice que el estudiante fue agregado y crea la nueva base de datos
    """
    base=open(archivo,"rb")
    lista=pickle.load(base)      #Descarga la lista de la base de datos actual.
    lista.append(estudiante)
    base.close()
    nuevaBase=open(archivo,"wb")    #Vuelve a crear la base de datos.
    pickle.dump(lista,nuevaBase)        #Añade la lista actualizada a la base de datos.
    nuevaBase.close()
    return "El estudiante ha sido agregado."

def agregarEstudianteAux(archivo,p1,p2,p3):
    """
    Funcionamiento: 
    - Solicita todos los datos del estudiante para llamar a la función principal de agregar estudiante.
    Entradas:
    - archivo(str): Es el archivo de la base de datos.
    - p1,p2,p3(int): Son los porcentajes que vale cada evaluación.
    Salidas: 
    - Llama a la función principal.
    """
    base=open(archivo,"rb")
    lista=pickle.load(base)
    while True:
        try:
            carne=input("Ingrese el carné del estudiante o digite 0 para cancelar:\n")        #Pide el número de carné del estudiante.
            if carne=="0":
                return ""
            elif not re.match(r"\d{10}$",carne):
                raise ValueError("El carné debe de tener 10 dígitos.")
            for i in lista:     #Verifica que el carné no se encuentre en la lista
                for j in i:        
                    if j==carne:
                        raise ValueError("El carné ingresado ya se encuentra registrado")
            break
        except ValueError as e:                        #Cada ValueError tiene su propio texto, el except toma el texto de el error que ocurre y lo imprime
            print(e)
    while True:
        try:
            nombre = tuple(input("Ingrese el nombre del estudiante con sus dos apellidos o digite 0 para cancelar:\n").split(" "))   #Pide el nombre del estudiante
            if nombre==("0"):
                return ""
            elif len(nombre)!=3:
                raise ValueError
            break
        except ValueError:
            print("Tiene que ingresar un nombre con dos apellidos.\n")
    while True:
        try:
            genero = input("Indique el género del estudiante:\n0.Cancelar\n1. Femenino\n2. Masculino\n")     #Pide el género del estudiante
            if genero not in ("1","2","0"):
                raise ValueError("Debe escoger una de las dos opciones.\n")
            elif genero=="0":
                return ""
            elif genero==1:
                genero=False
            else:
                genero=True
            break
        except:
            print("Debe escoger una de las dos opciones.\n")
    while True:
        try:
            correo=input("Ingrese el correo del estudiante o digite 0 para cancelar")     #Pide el correo del estudiante
            if correo=="0":
                return ""
            if not re.match(r"\w+@estudiantec.cr$",correo):
                raise ValueError("El correo debe terminar en @estudiantec.cr")
            for i in lista:   #Verifica que el correo no esté en la lista
                for j in i:
                    if j==correo:
                        raise ValueError("El correo ingresado ya se encuentra ocupado por otra persona, inténtelo de nuevo.\n")
            break
        except ValueError as e:     
            print(e)
    while True:
        try:
            x1=int(input("Ingrese la nota obtenida por el estudiante en el primer rubro o digite 101 para cancelar: "))     #Pide las 3 notas del estudiante
            x2=int(input("Ingrese la nota obtenida por el estudiante en el segundo rubro o digite 101 para cancelar: "))
            x3=int(input("Ingrese la nota obtenida por el estudiante en el tercer rubro o digite 101 para cancelar: "))
            for i in (x1,x2,x3):
                if i==101:
                    return ""
                elif i>100 or i<0:
                    raise ValueError
            prom=round((x1*p1/100)+(x2*p2/100)+(x3*p3/100),2)       #Calcula el promedio de notas
            nota=(x1,x2,x3,prom,prom)
            break
        except ValueError:
            print("Las notas deben estar entre 0 y 100.\n")
    estudiante = [nombre,genero,carne,correo,nota]   #Son todos los datos del estudiante
    base.close()
    return agregarEstudiante(archivo,estudiante)

def obtenerNota(i):
    """
    Funcionamiento: 
    - Se utiliza para obtener la nota del estudiante desde su lista de datos, se utiliza en la función de reporteGenero para ordenar los nombres según la nota.
    Entradas:
    - i(list): Es la lista que contiene los datos del estudiante.
    Salidas: 
    - Retorna la nota del estudiante.
    """
    return i[4][4] 

def reporteGenero(archivo,x1,x2,x3):
    """
    Funcionamiento: 
    - Se utiliza para generar los archivos de reporte por género.
    Entradas:
    - archivo(str): Es el archivo de la base de datos.
    - x1,x2,x3(int): Son los porcentajes que vale cada rubro.
    Salidas: 
    - Crea los archivos con los reportes e imprime un texto que dice que los archivos fueron agregados a la carpeta.
    """
    contadorHombres=0
    contadorMujeres=0
    mujeres = Document()    #Crea un documento y se le asigna a la variable mujeres.
    mujeres.add_heading("Reporte de Notas Mujeres", level=1)     #Añade el título.
    mujeres.add_paragraph("")           #Deja un espacio en blanco
    hombres = Document()
    hombres.add_heading("Reporte de Notas Hombres", level=1)
    hombres.add_paragraph("")
    base=open(archivo,"rb")
    personas = pickle.load(base)             #Es la lista que contiene los datos de todas las personas.
    personas.sort(key=obtenerNota,reverse=True)     #La lista se ordena de acuerdo a la nota dada por la función obtenerNota de manera inversa, o sea, de mayor a menor.
    for i in personas:
        linea=f"{str(i[4][4])}, {str(i[4][0])} {str(i[4][1])} {str(i[4][2])}, {i[0][0]} {i[0][1]} {i[0][2]}, {i[2]}, {i[3]}"   #Es la linea que se agrega al archivo creado con la información de la persona.
        if i[1]==False:     #Si el género el False se agrega a mujeres
            contadorMujeres+=1
            mujeres.add_paragraph(linea)
        else:                            #De lo contrario se agrega a hombres
            contadorHombres+=1
            hombres.add_paragraph(linea)
    mujeres.add_paragraph(f"\nLos porcentajes de cada evaluación fueron {x1}%, {x2}% y {x3}% respectivamente, y la cantidad de mujeres es {contadorMujeres}.")   #Agrega una línea que dice los porcentajes originales y la totalidad de personas en el documento.
    hombres.add_paragraph(f"\nLos porcentajes de cada evaluación fueron {x1}%, {x2}% y {x3}% respectivamente, y la cantidad de hombres es {contadorHombres}.")
    base.close()
    mujeres.save("mujeres.docx")   #Se guarda el archivo con el nombre de mujeres.docx
    hombres.save("hombres.docx")
    return print("Los archivos fueron agregados a la carpeta")

def reporteGeneracion(archivo):
    """
    Funcionamiento: 
    - Genera un reporte por generación que separa a las personas según su nota.
    Entradas:
    - archivo(str): Es el archivo que contiene la base de datos.
    Salidas: 
    - Imprime el reporte.
    """
    apTotales=0
    rpTotales=0
    reTotales=0
    gens=[]
    gensFinal=[]
    base=open(archivo,"rb")
    lista=pickle.load(base)
    for i in lista:
        if i[2][:4] not in gens:    #Lee las generaciones de los carnés y las agrega a la lista gens.
            gens.append(int(i[2][:4]))
    gens.sort()                   
    for i in gens:
        gensFinal.append([0,0,0])     #Según la cantidad de generaciones, se agrega una lista con tres ceros a la lista gensFinal
    for i in lista:
        if i[4][4]>=70:   #Si la persona aprobó, en gensFinal se le suma uno al primer 0 que se encuentra en el índice en que está la generación del estudiante en gens.
            gensFinal[gens.index(int(i[2][:4]))][0]+=1
        elif 60<=i[4][4]<70:
            gensFinal[gens.index(int(i[2][:4]))][1]+=1
        else:
            gensFinal[gens.index(int(i[2][:4]))][2]+=1
    print("Generación\tAprobados\tReposición\tRebrobados\tTotales")
    for i in range(len(gens)):
        print(f"   {gens[i]}\t\t    {gensFinal[i][0]}\t\t    {gensFinal[i][1]}\t\t    {gensFinal[i][2]}\t\t   {gensFinal[i][0]+gensFinal[i][1]+gensFinal[i][2]}")    #Se imprime el reporte de cada sede con su total de aprobados,aplazados y reprobados.
        apTotales+=gensFinal[i][0]
        rpTotales+=gensFinal[i][1]    #Los aprobados, aplazados y reprobados se van sumando a un reporte total.
        reTotales+=gensFinal[i][2]
    total = apTotales+rpTotales+reTotales
    print(f"  Totales\t    {apTotales}\t\t    {rpTotales}\t\t    {reTotales}\t\t   {total}")      #Se imprime un reporte total
    base.close()
    return""

def curvasHtml(archivo, lista):
    """
    Funcionamiento:
    - Es la auxiliar que solicita el porcentaje de curva que se le va a aplicar a la nota final.
    Entradas:
    archivo: contiene la información de la base de datos
    lista: es la lista en la que se trabaja la base de datos
    Salidas:
    Se retorna el valor de salida de la función curvaAprovado.
    """
    while True:
        try:
            os.system("cls")
            print("******************** Gestionar Curva ********************")
            porcentaje=int(input("Ingrese el porcentaje de curva a aplicar: "))
            os.system("cls")
            if not 0<porcentaje<100:
                raise TypeError
            return curvaAprovado(archivo, lista, porcentaje)
        except TypeError:
            print("******************** Gestionar Curva ********************")
            print("El porcentaje debe ser mayor a 0 y menor que 100.")
            input("Presione enter para continuar.")
            os.system("cls")
        except ValueError:
            os.system("cls")
            print("******************** Gestionar Curva ********************")
            print("Debe ingresar valores válidos. Por ejemplo:\n" \
            "Indique el porcentaje de curva a aplicar: 5")
            input("Presione enter para continuar.")
            os.system("cls")

def enviarCorreos(archivo,fecha,hora):
    """
    Funcionamiento: 
    - Se utiliza para enviar los correos con la fecha del examen de reposición.
    Entradas:
    - archivo(str): Es el archivo de la base de datos.
    - fecha(str): Es la fecha en la que se va a realizar el examen.
    - hora(str): Es la hora a la que se realizará el examen.
    Salidas: 
    - Envía los correos y retorna un mensaje indicándolo
    """
    base=open(archivo,"rb")
    lista=pickle.load(base)
    correoRemitente = "xaviccr@gmail.com"   #Es el correo desde el cual se van a enviar los correos.
    contrasenna = "tqge uhju oxon ibgu" #Es la contraseña que se utiliza para ingresar al correo.
    for i in lista:
        if 60<=i[4][4]<70:
            correoEstudiante = i[3]
            asunto = "Examen de reposición"
            mensaje = f"Se le comunica que usted deberá realizar un examen de reposición el día {fecha} a las {hora}"
            msg = MIMEMultipart()     #Sirve para dar formato al mensaje con el correo remitente, el correo receptor y el asunto.
            msg["From"] = correoRemitente
            msg["To"] = correoEstudiante
            msg["Subject"] = asunto
            msg.attach(MIMEText(mensaje, 'plain'))    # a msg se le añade el mensaje en texto plano para darle forma al correo que se va a enviar.
            servidor = smtplib.SMTP('smtp.gmail.com', 587) #Entra al servidor que permite enviar los correos, 587 es el puerto que se utiliza para esta acción.
            servidor.starttls()    #Crea una capa de seguridad para proteger los datos.
            servidor.login(correoRemitente,contrasenna)  #Entra al correo remitente para enviar los correos desde allí
            servidor.send_message(msg)   #Envía el correo
            servidor.quit()    #Sale del servidor
    base.close()
    return "Los correos para reposición han sido enviados"

def enviarCorreosAux(archivo):
    """
    Funcionamiento: 
    - Se utiliza para solicitar la fecha y la hora de los exámenes de reposición.
    Entradas:
    - archivo(str): Es el archivo de la base de datos.
    Salidas: 
    - Llama a la función principal de enviar correos.
    """
    while True:
        try:
            fecha=input("Ingrese el día en que se va a realizar el examen de reposición con el formato dd/mm/aaaa:\n")
            if not re.match(r"^((0[1-9]|1\d|2\d)/(0[1-9]|1[0-2])|(30)/(0[1,3-9]|1[0-2])|31/(0[13578]|1[02]))/(\d{4})$",fecha):   #Valida si la fecha cumple con el formato.
                raise ValueError("La fecha no cumple con el formato requerido de dd/mm/aaaa, recuerde que el mes no puede ser mayor a 12 y que solo ciertos meses tienen 31 días.\n")
            elif int(fecha[0:2]) == 29 and int(fecha[3:5]) == 2:
                if not ((int(fecha[6:]) % 4 == 0 and int(fecha[6:]) % 100 != 0) or (int(fecha[6:]) % 400 == 0)):   #Si la fecha ingresada es el 29/02/aaaa, verifica si el año es bisiesto.
                    raise ValueError("\nLa fecha no es válida, pues el año ingresado no es bisiesto\n")
            dia, mes, anno = int(fecha[:2]), int(fecha[3:5]), int(fecha[6:])
            if datetime(anno, mes, dia)<=datetime.now():   #Verifica si la fecha es posterior a la actual.
                raise ValueError("Debe ingresar una fecha posterior a la de hoy.\n")
            break
        except ValueError as e:
            print(e)
    while True:
        try:
            hora=input("Ingrese la hora en un formato de 24 horas, o sea, hh:mm:\n")
            if not re.match(r"^([01]\d|2[0-3]):[0-5]\d$",hora):   #Verifica el formato de la hora.
                raise ValueError
            break
        except ValueError:
            print("La hora no concuerda con el formato de 24 horas\n")
    return enviarCorreos(archivo,fecha,hora)

def reporteBuenRendimiento(archivo,sedes):
    """
    Funcionamiento:
    - Pide al usuario ingresar un número de sede e imprime los estudiantes de esa sede que obtuvieron notas mayores a 70.
    Entradas:
    - archivo: Es el archivo de la base de datos.
    - sedes: Es el archivo que contiene las sedes.
    Salidas: 
    - Imprime el nombre y el carné de los estudiantes.
    """
    codigos=[]
    estudiantes=[]
    base=open(archivo,"rb")
    lista=pickle.load(base)
    sedes=open(sedes,"r",encoding="utf-8")   #Se usa encoding="utf-8" para leer correctamente las líneas del archivo, ya que estas se van a imprimir.
    contador=0
    for i,linea in enumerate(sedes):   #Enumera las sedes y las imprime.
        print(f"Sede {i+1}, {linea}")    
        if len(str(i))==1 and i!=9:      #Si se encuentra en las líneas del 0-8, se le agrega un cero a la izquierda para que el código concuerde con los carnés.
            codigos.append(f"0{i+1}")
        else: 
            codigos.append(i+1)
    while True:
        try:
            sede=input("\nIngrese el número de la sede de la que desea generar el reporte:\n")
            if len(sede)==1:
                sede="0"+sede
            if sede not in codigos:   #Verifica si la sede ingresada por el usuario se encuentra en la lista de sedes.
                raise ValueError
            break
        except ValueError:
            print("Tiene que digitar el código de una de las sedes antes mostradas.\n")
    for i in lista:
        if i[2][4:6] == sede:
            if i[4][0]>=70 and i[4][1]>=70 and i[4][2]>=70:
                contador+=1             #Enumera los estudiantes
                estudiante=f"{contador}. {i[0][0]} {i[0][1]} {i[0][2]}, {i[2]}"     #agrega el nombre y el carné del estudiante a estudiantes.
                estudiantes.append(estudiante)
    if len(estudiantes)==0:     #Detecta si no hubo estudiantes con buen rendimiento.
        print("\nNo existen estudiantes con buen rendimiento en dicha sede.\n")
    else:
        print("\nLos estudiantes que demostraron un buen rendimiento en la sede solicitada fueron:\n")
        for i in estudiantes:     #Imprime los datos de cada estudiante.
            print(i)
    base.close()
    return ""